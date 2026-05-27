"""Verify a translated DOCX/PDF for mid-paragraph English leaks.

This is the diagnostic that was missing the night the saathi curriculum doc shipped
with a Hindi paragraph followed by an untranslated red English chunk: the original
verification only flagged paragraphs that were ENTIRELY English, missing paragraphs that
were mostly Hindi but had a leaked English tail.

A leak is: a paragraph that contains Devanagari/Indic script AND a substantial run of
consecutive English alphabetic words (>= MIN_LEAK_WORDS). URLs, email addresses, proper
nouns (single capitalised words), and short acronyms (≤4 letters) are not counted.

Usage:
    python verify_output.py /tmp/translated.docx
    python verify_output.py /tmp/translated.docx --target hi --min-words 4 --strict

Exit codes:
    0  no leaks found
    1  leaks found (use --strict to fail CI)
    2  could not parse the file
"""
import argparse
import re
import sys
from pathlib import Path

# Each tuple is (lang_code, regex matching at least one char in that script).
_INDIC_SCRIPTS = [
    ("hi/mr", re.compile(r"[ऀ-ॿ]")),
    ("bn/as", re.compile(r"[ঀ-৿]")),
    ("pa",    re.compile(r"[਀-੿]")),
    ("gu",    re.compile(r"[઀-૿]")),
    ("or",    re.compile(r"[଀-୿]")),
    ("ta",    re.compile(r"[஀-௿]")),
    ("te",    re.compile(r"[ఀ-౿]")),
    ("kn",    re.compile(r"[ಀ-೿]")),
    ("ml",    re.compile(r"[ഀ-ൿ]")),
    ("ur",    re.compile(r"[؀-ۿ]")),
]

# A "real" English word: 2+ letters, mostly lowercase or sentence-cased. We exclude:
# - URLs/emails (split on whitespace, then a word containing :// or @)
# - ALL-CAPS short acronyms (AI, NGO, GST) — likely intentional
# - Numbers / mixed alphanumeric (page 12, IPv4) — not text content
_URL_RE   = re.compile(r"https?://\S+|www\.\S+|\S+@\S+\.\S+")
_WORD_RE  = re.compile(r"[A-Za-z][A-Za-z\-']{1,}")
_ACRO_RE  = re.compile(r"^[A-Z]{2,4}$")
_PROP_RE  = re.compile(r"^[A-Z][a-z]+$")  # single capitalised — likely a name


def has_indic(text: str) -> str | None:
    for name, rx in _INDIC_SCRIPTS:
        if rx.search(text):
            return name
    return None


def english_words(text: str) -> list[str]:
    """Return the list of English-alphabetic words that look like prose (not URL / acronym /
    single proper noun). Each occurrence stays in source order."""
    cleaned = _URL_RE.sub(" ", text)
    out = []
    for w in _WORD_RE.findall(cleaned):
        if _ACRO_RE.match(w):
            continue
        if _PROP_RE.match(w) and len(w) > 3:  # multi-letter capitalised: probably proper noun
            continue
        out.append(w)
    return out


def longest_english_run(text: str) -> int:
    """Approximate the size of the LARGEST consecutive English run in the paragraph (in
    word count). A mixed paragraph with one short English phrase rates low; a paragraph
    with an entire untranslated English sentence rates high."""
    # Tokenise by whitespace; mark each token as English-content / not.
    cleaned = _URL_RE.sub(" ", text)
    toks = cleaned.split()
    streaks = []
    cur = 0
    for t in toks:
        # strip surrounding punctuation
        bare = re.sub(r"^[^A-Za-z]+|[^A-Za-z]+$", "", t)
        if _WORD_RE.fullmatch(bare) and len(bare) >= 2 and not _ACRO_RE.match(bare) and not _PROP_RE.match(bare):
            cur += 1
        else:
            if cur:
                streaks.append(cur)
            cur = 0
    if cur:
        streaks.append(cur)
    return max(streaks) if streaks else 0


def scan_docx(path: Path):
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(path))
    paras = list(doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    for section in doc.sections:
        for hf in (section.header, section.footer):
            paras.extend(hf.paragraphs)

    w_t = qn("w:t")
    out = []
    for i, p in enumerate(paras):
        text = "".join(t.text or "" for t in p._p.iter(w_t))
        if text.strip():
            out.append((i, text))
    return out


def scan_pdf(path: Path):
    import fitz

    doc = fitz.open(str(path))
    out = []
    for pi, page in enumerate(doc):
        for bi, b in enumerate(page.get_text("dict")["blocks"]):
            if b.get("type") != 0:
                continue
            text = " ".join(s["text"] for line in b["lines"] for s in line["spans"]).strip()
            if text:
                out.append((f"p{pi + 1}.b{bi + 1}", text))
    return out


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.split("\n\n")[0])
    ap.add_argument("path", help="translated .docx or .pdf to scan")
    ap.add_argument("--target", default=None,
                    help="expected target language code (hi/mr/...) — if set, only paragraphs that "
                         "contain THAT script count as candidates for the mixed-leak check")
    ap.add_argument("--min-words", type=int, default=3,
                    help="report paragraphs whose longest English run is at least this many words")
    ap.add_argument("--show", type=int, default=10, help="number of offending paragraphs to print")
    ap.add_argument("--strict", action="store_true", help="exit 1 if any leak found")
    args = ap.parse_args()

    p = Path(args.path)
    if not p.exists():
        print(f"ERROR: file not found: {p}", file=sys.stderr)
        return 2

    ext = p.suffix.lower()
    if ext == ".docx":
        items = scan_docx(p)
    elif ext == ".pdf":
        items = scan_pdf(p)
    else:
        print(f"ERROR: unsupported extension: {ext} (need .docx or .pdf)", file=sys.stderr)
        return 2

    full_english = []   # paragraphs that contain NO target script at all
    mixed_leaks = []    # paragraphs with target script AND a substantial English run
    indic_paragraphs = 0
    for ident, text in items:
        script = has_indic(text)
        if script:
            indic_paragraphs += 1
            run = longest_english_run(text)
            if run >= args.min_words:
                if args.target is None or script.split("/")[0] == args.target or args.target in script:
                    mixed_leaks.append((ident, run, text))
        else:
            if english_words(text):
                full_english.append((ident, text))

    print(f"Scanned {len(items)} paragraphs / blocks ({indic_paragraphs} with Indic script)")
    print(f"  fully-English (untranslated):  {len(full_english)}")
    print(f"  MIXED leaks (>= {args.min_words} English words after Indic): {len(mixed_leaks)}")

    if mixed_leaks:
        print()
        print(f"Top {min(args.show, len(mixed_leaks))} mixed-leak paragraphs:")
        for ident, run, text in sorted(mixed_leaks, key=lambda x: -x[1])[: args.show]:
            preview = text if len(text) <= 240 else text[:120] + " […] " + text[-120:]
            print(f"  • [{ident}]  longest-English-run={run} words")
            print(f"      {preview}")

    if full_english and args.show:
        print()
        print(f"Top {min(args.show, len(full_english))} fully-English paragraphs:")
        for ident, text in full_english[: args.show]:
            preview = text if len(text) <= 200 else text[:200] + "…"
            print(f"  • [{ident}]  {preview}")

    if args.strict and (mixed_leaks or full_english):
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
